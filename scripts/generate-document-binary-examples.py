#!/usr/bin/env python3
"""Generate independent PDF, ePub, and DOCX fixtures for document translation tests."""

from __future__ import annotations

import argparse
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer


def set_run_font(run, name: str, size: float, color: str = "202533", bold: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def generate_docx(output: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, "2E74B5", 18, 10),
        "Heading 2": (13, "2E74B5", 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("FLUENTREAD · DOCUMENT TRANSLATION EXAMPLE"), "Calibri", 8.5, "7A8294", True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    set_run_font(footer.add_run("Local browser fixture"), "Calibri", 8.5, "7A8294")

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(6)
    set_run_font(kicker.add_run("REFERENCE GUIDE"), "Calibri", 10, "E83B6B", True)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    set_run_font(title.add_run("Document Translation Example"), "Calibri", 26, "202533", True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(18)
    set_run_font(subtitle.add_run("A realistic DOCX fixture for bilingual reading regression"), "Calibri", 13, "5F687C")

    document.add_heading("Why local document translation matters", level=1)
    document.add_paragraph(
        "A useful document translator should preserve the source file while presenting each translated paragraph directly beneath its original text."
    )
    paragraph = document.add_paragraph()
    set_run_font(paragraph.add_run("Privacy first. "), "Calibri", 11, "202533", True)
    set_run_font(
        paragraph.add_run("The file is parsed inside the browser, and only extracted text is sent to the translation service selected by the reader."),
        "Calibri",
        11,
    )

    document.add_heading("What the regression should protect", level=1)
    for label, detail in (
        ("Structure", "Headings, paragraph order, headers, and footers remain valid after export."),
        ("Reading flow", "The original paragraph appears first and the editable translation follows naturally."),
        ("Format", "The downloaded result remains a standards-compliant DOCX that opens in Word and LibreOffice."),
    ):
        paragraph = document.add_paragraph()
        set_run_font(paragraph.add_run(f"{label}: "), "Calibri", 11, "2E74B5", True)
        set_run_font(paragraph.add_run(detail), "Calibri", 11)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def generate_pdf(output: Path) -> None:
    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "FixtureTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=24,
        leading=29,
        textColor=HexColor("#202533"),
        spaceAfter=16,
    )
    heading = ParagraphStyle(
        "FixtureHeading",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=20,
        textColor=HexColor("#E83B6B"),
        spaceBefore=10,
        spaceAfter=10,
    )
    body = ParagraphStyle(
        "FixtureBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=11,
        leading=16,
        textColor=HexColor("#202533"),
        spaceAfter=10,
    )
    story = [
        Paragraph("Document Translation Example", title),
        Paragraph("A text-layer PDF used to verify page-aware bilingual reading and export.", body),
        Spacer(1, 0.12 * inch),
        Paragraph("Readable PDF text", heading),
        Paragraph(
            "FluentRead extracts text from each PDF page locally, groups it into readable segments, and keeps a page label beside the first segment.",
            body,
        ),
        Paragraph(
            "The bilingual download keeps the original page and follows it with a generated translation page, so the source remains available for comparison.",
            body,
        ),
        PageBreak(),
        Paragraph("Regression coverage", title),
        Paragraph("The second page proves that page boundaries survive parsing and remain visible in the reading interface.", body),
        Paragraph(
            "A scanned PDF without a text layer is rejected with a clear OCR-specific message instead of producing empty or corrupted output.",
            body,
        ),
    ]
    output.parent.mkdir(parents=True, exist_ok=True)
    pdf = SimpleDocTemplate(
        str(output),
        pagesize=letter,
        leftMargin=0.85 * inch,
        rightMargin=0.85 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        title="Document Translation Example",
        author="FluentRead",
    )
    pdf.build(story)


def generate_epub(output: Path) -> None:
    mimetype = "application/epub+zip"
    container = """<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles><rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/></rootfiles>
</container>
"""
    content_opf = """<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="book-id">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:identifier id="book-id">fluentread-document-example</dc:identifier>
    <dc:title>Document Translation Example</dc:title>
    <dc:language>en</dc:language>
  </metadata>
  <manifest>
    <item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>
    <item id="chapter-1" href="chapter-1.xhtml" media-type="application/xhtml+xml"/>
    <item id="chapter-2" href="chapter-2.xhtml" media-type="application/xhtml+xml"/>
  </manifest>
  <spine>
    <itemref idref="chapter-1"/>
    <itemref idref="chapter-2"/>
  </spine>
</package>
"""
    nav = """<?xml version="1.0" encoding="UTF-8"?>
<html xmlns="http://www.w3.org/1999/xhtml"><head><title>Contents</title></head>
<body><nav epub:type="toc" xmlns:epub="http://www.idpf.org/2007/ops"><ol>
<li><a href="chapter-1.xhtml">Fluent reading</a></li><li><a href="chapter-2.xhtml">Regression</a></li>
</ol></nav></body></html>
"""
    chapter_1 = """<?xml version="1.0" encoding="UTF-8"?>
<html xmlns="http://www.w3.org/1999/xhtml"><head><title>Fluent reading</title></head><body>
<h1>Fluent reading for local books</h1>
<p>An ePub translator should preserve chapter order, links, and the original source while placing each translation directly below its paragraph.</p>
<p>Readers can edit the translation before downloading a bilingual electronic book.</p>
</body></html>
"""
    chapter_2 = """<?xml version="1.0" encoding="UTF-8"?>
<html xmlns="http://www.w3.org/1999/xhtml"><head><title>Regression coverage</title></head><body>
<h1>Regression coverage</h1>
<p>The exported archive must remain a valid ePub package with an uncompressed mimetype entry and readable XHTML chapters.</p>
<p><a href="chapter-1.xhtml">Return to the first chapter</a> without changing the link target.</p>
</body></html>
"""

    output.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output, "w") as archive:
        archive.writestr("mimetype", mimetype, compress_type=zipfile.ZIP_STORED)
        archive.writestr("META-INF/container.xml", container, compress_type=zipfile.ZIP_DEFLATED)
        archive.writestr("OEBPS/content.opf", content_opf, compress_type=zipfile.ZIP_DEFLATED)
        archive.writestr("OEBPS/nav.xhtml", nav, compress_type=zipfile.ZIP_DEFLATED)
        archive.writestr("OEBPS/chapter-1.xhtml", chapter_1, compress_type=zipfile.ZIP_DEFLATED)
        archive.writestr("OEBPS/chapter-2.xhtml", chapter_2, compress_type=zipfile.ZIP_DEFLATED)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    generate_pdf(args.output_dir / "sample.pdf")
    generate_epub(args.output_dir / "sample.epub")
    generate_docx(args.output_dir / "sample.docx")


if __name__ == "__main__":
    main()
